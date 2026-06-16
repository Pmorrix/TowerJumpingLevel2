# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\TowerJumpingLevel2\Guion_Clase_Practica_Cube_Runner_Arena_IA_COMPLETO.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "596579"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PROMPT_FILL = "F7F9FC"
RED = "D71920"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D6DEE8", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, dxa=9360, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=None, bold=None, color=None, font="Calibri", italic=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.25, left_indent=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if left_indent is not None:
        paragraph.paragraph_format.left_indent = Inches(left_indent)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=18 if level == 1 else 12, after=8 if level == 1 else 6)
    run = p.add_run(text)
    if level == 1:
        style_run(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        style_run(run, size=13, bold=True, color=BLUE)
    else:
        style_run(run, size=12, bold=True, color=DARK_BLUE)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    set_paragraph_format(p, after=after)
    run = p.add_run(text)
    style_run(run, size=11, color=INK)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, after=4, line_spacing=1.25)
        run = p.add_run(item)
        style_run(run, size=11, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_format(p, after=4, line_spacing=1.25)
        run = p.add_run(item)
        style_run(run, size=11, color=INK)


def add_callout(doc, title, lines, fill=LIGHT_GRAY, accent=BLUE, mono=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="C9D3E0", size="8")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    set_paragraph_format(p, after=4)
    r = p.add_run(title)
    style_run(r, size=10.5, bold=True, color=accent)
    font = "Consolas" if mono else "Calibri"
    if isinstance(lines, str):
        lines = lines.strip("\n").split("\n")
    for line in lines:
        p = cell.add_paragraph()
        set_paragraph_format(p, after=2, line_spacing=1.15)
        r = p.add_run(line)
        style_run(r, size=9.5 if mono else 10.5, color=INK, font=font)
    doc.add_paragraph()


def add_schedule_table(doc):
    data = [
        ("0:00-0:15", "Apertura", "Presentar el reto y explicar que se comparara el razonamiento humano con la solucion de IA."),
        ("0:15-1:15", "Terreno de juego", "Construir la escena con primitivas basicas y preparar Player, coleccionables, obstaculos, camara, UI y GameManager."),
        ("1:15-1:30", "Puesta en comun", "Revisar errores tipicos de escena, componentes, tags, colliders e Inspector."),
        ("1:30-2:20", "Diseno manual", "Definir scripts, responsabilidades, GameObjects y referencias antes de usar IA."),
        ("2:20-2:30", "Descanso", "Pausa breve."),
        ("2:30-3:20", "Desarrollo con ChatGPT", "Pedir a la IA primero la arquitectura y despues el codigo completo."),
        ("3:20-4:10", "Integracion y debugging", "Copiar scripts, asignar Inspector, probar Play Mode y depurar errores."),
        ("4:10-4:40", "Comparacion", "Comparar propuesta humana, propuesta de IA y solucion final corregida."),
        ("4:40-5:00", "Cierre", "Entrega minima, conclusiones y mensaje final."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    headers = ["Tiempo", "Bloque", "Objetivo"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        style_run(r, bold=True, color=DARK_BLUE, size=10.5)
    set_repeat_table_header(table.rows[0])
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_borders(cells[i])
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            set_paragraph_format(p, after=0, line_spacing=1.15)
            r = p.add_run(text)
            style_run(r, size=9.5 if i == 0 else 10, color=INK)
    doc.add_paragraph()


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for style_name in ("List Bullet", "List Number"):
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(11)
    return doc


def build():
    doc = setup_doc()

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Master en Creacion de Videojuegos - Universidad de Malaga")
    style_run(r, size=9, color=MUTED)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=80, after=6)
    r = p.add_run("Cube Runner Arena")
    style_run(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=4)
    r = p.add_run("Guion completo de clase practica con IA")
    style_run(r, size=18, color=INK)

    p = doc.add_paragraph()
    set_paragraph_format(p, after=18)
    r = p.add_run("Uso de ChatGPT gratuito en Unity 3D y C#")
    style_run(r, size=13, color=MUTED)

    add_callout(
        doc,
        "Resumen de la sesion",
        [
            "Duracion: 5 horas",
            "Bloque: Inteligencia Artificial aplicada al desarrollo de videojuegos",
            "Herramienta de apoyo: ChatGPT gratuito, sin suscripcion",
            "Producto final: minijuego jugable, scripts integrados y comparacion entre solucion humana e IA.",
        ],
        fill="F1F3F5",
        accent=RED,
    )
    doc.add_page_break()

    add_heading(doc, "Objetivo general", 1)
    add_body(doc, "Crear un minijuego completo en Unity usando primitivas basicas y apoyandose en ChatGPT como asistente tecnico.")
    add_body(doc, "El alumnado primero razonara la estructura del minijuego y los scripts necesarios. Despues pedira a la IA que proponga la arquitectura y genere los scripts. Finalmente se compararan ambas soluciones.")
    add_body(doc, "La idea principal de la clase no es solo que el juego funcione, sino aprender a evaluar criticamente lo que propone una IA.")

    add_heading(doc, "Minijuego propuesto", 1)
    add_body(doc, "Nombre: Cube Runner Arena")
    add_body(doc, "El jugador controla una capsula dentro de una arena sencilla. Debe recoger todas las esferas antes de que termine el tiempo y evitar los cubos rojos en movimiento.")
    add_heading(doc, "Elementos de la escena", 2)
    add_bullets(doc, [
        "Player: una capsula.",
        "Coleccionables: esferas.",
        "Obstaculos: cubos rojos.",
        "Suelo: un plano.",
        "Paredes: cubos.",
        "Camara: fija, mostrando toda la arena.",
        "UI: puntuacion, tiempo y mensaje final.",
    ])
    add_heading(doc, "Condiciones de juego", 2)
    add_bullets(doc, [
        "El jugador se mueve por el plano XZ.",
        "Cada esfera recogida suma puntos.",
        "Al recoger todas las esferas, el jugador gana.",
        "Si el jugador toca un obstaculo, pierde.",
        "Si el tiempo llega a 0, pierde.",
        "Al aparecer Game Over, el Player se desactiva.",
        "La escena se puede reiniciar pulsando R.",
    ])
    add_callout(doc, "Resultado esperado al final de la practica", [
        "Escena jugable.",
        "Scripts funcionando.",
        "Game Over.",
        "Victoria.",
        "Reinicio.",
        "Comparacion entre razonamiento humano e IA.",
    ])

    add_heading(doc, "Estructura temporal", 1)
    add_schedule_table(doc)

    add_heading(doc, "0:00 - 0:15 | Apertura de la practica", 1)
    add_body(doc, "El profesor presenta el reto y situa la practica dentro del bloque de IA.")
    add_callout(doc, "Texto para el profesor", "Vamos a crear un minijuego muy pequeno, pero completo. Primero pensaremos como desarrolladores: que objetos necesita la escena, que comportamientos hacen falta y que scripts deberian existir. Despues pediremos a ChatGPT que proponga los scripts y compararemos su solucion con la nuestra.")
    add_heading(doc, "Ideas clave", 2)
    add_bullets(doc, [
        "La IA no sustituye la comprension de Unity.",
        "ChatGPT puede escribir codigo, pero no ve automaticamente la escena.",
        "En Unity, el comportamiento final depende de codigo, GameObjects, componentes, Inspector, fisicas, tags y UI.",
        "Una buena solucion no es la mas grande, sino la mas clara y mantenible.",
    ])
    add_callout(doc, "Objetivo de esta primera parte", [
        "Que objetos hay.",
        "Que comportamientos hacen falta.",
        "Que eventos ocurren.",
        "Que datos se comparten.",
        "Que referencias van por Inspector.",
        "Que condiciones hacen ganar o perder.",
    ])

    add_heading(doc, "0:15 - 1:15 | Preparar el terreno de juego", 1)
    add_body(doc, "Durante la primera hora, el alumnado prepara manualmente la escena de Unity.")
    add_heading(doc, "Tarea del alumnado", 2)
    add_bullets(doc, [
        "Un Plane como suelo.",
        "Una Capsule como Player.",
        "Un Rigidbody en el Player.",
        "Tag Player asignado al Player.",
        "Cubos como paredes.",
        "Esferas como coleccionables.",
        "Cubos rojos como obstaculos.",
        "Camara fija.",
        "Luz principal.",
        "Canvas con textos de UI.",
        "Un GameObject vacio llamado GameManager.",
    ])
    add_callout(doc, "UI necesaria", [
        "ScoreText",
        "TimeText",
        "MessageText",
        "MessageText puede empezar vacio o desactivado.",
    ])
    add_callout(doc, "Checklist de escena", [
        "Player tiene Rigidbody.",
        "Player tiene Collider.",
        "Player tiene tag Player.",
        "Coleccionables tienen Collider con Is Trigger.",
        "Obstaculos tienen Collider normal.",
        "Camara ve toda la arena.",
        "UI existe.",
        "GameManager esta en escena.",
    ])
    add_heading(doc, "Si no terminan la escena", 2)
    add_body(doc, "Si pasado este tiempo algun alumno o grupo no ha conseguido preparar el terreno, se entregara un .unitypackage con la escena ya montada, pero sin scripts.")
    add_callout(doc, "Continuacion comun para todo el grupo", [
        "Pensar los scripts.",
        "Pedir ayuda a la IA.",
        "Integrar codigo.",
        "Depurar.",
        "Comparar soluciones.",
    ])

    add_heading(doc, "1:15 - 1:30 | Puesta en comun de la escena", 1)
    add_heading(doc, "Revision rapida", 2)
    add_bullets(doc, [
        "Falta el tag Player.",
        "El Player no tiene Rigidbody.",
        "Los coleccionables no tienen Is Trigger.",
        "Los obstaculos tienen Is Trigger por error.",
        "La camara no ve bien la arena.",
        "La UI no esta enlazada.",
        "Falta el objeto GameManager.",
    ])
    add_callout(doc, "Mensaje clave", "En Unity, gran parte del comportamiento depende de la escena y del Inspector, no solo del codigo.")
    add_callout(doc, "Preguntas al alumnado", [
        "Que objeto deberia detectar que se ha recogido una esfera?",
        "Quien deberia guardar la puntuacion?",
        "Quien deberia controlar el tiempo?",
        "Quien decide si se gana o se pierde?",
        "Donde deberia ir el codigo de movimiento?",
        "Que referencias se asignan desde Inspector?",
    ])

    add_heading(doc, "1:30 - 2:20 | Diseno manual de los scripts", 1)
    add_body(doc, "Antes de usar ChatGPT, el alumnado debe proponer que scripts cree que necesita el juego.")
    add_callout(doc, "Actividad", [
        "Lista de scripts necesarios.",
        "Responsabilidad de cada script.",
        "GameObject donde iria colocado cada script.",
        "Referencias necesarias en Inspector.",
        "Eventos importantes.",
    ])
    add_callout(doc, "Preguntas guia", [
        "Que script debe mover al Player?",
        "Que script debe detectar una esfera recogida?",
        "Que script debe controlar el score?",
        "Que script debe controlar el tiempo?",
        "Que script debe decidir victoria o derrota?",
        "Que script debe mover los obstaculos?",
        "Que script debe detectar choque con obstaculo?",
        "Que referencias deben ir por Inspector?",
    ])
    add_callout(doc, "Solucion minima esperada", [
        "PlayerMovement.cs",
        "Collectible.cs",
        "GameManager.cs",
        "MovingObstacle.cs",
        "",
        "No se da esta lista como respuesta cerrada al principio. Primero se deja que el alumnado razone.",
    ])
    add_heading(doc, "Responsabilidad de cada script", 2)
    add_callout(doc, "PlayerMovement.cs", [
        "Mover el Player con teclado.",
        "Usar Rigidbody.",
        "Movimiento en plano XZ.",
        "No decidir victoria ni derrota.",
    ])
    add_callout(doc, "Collectible.cs", [
        "Detectar al Player con OnTriggerEnter.",
        "Avisar al GameManager.",
        "Sumar puntos.",
        "Desactivar la esfera recogida.",
    ])
    add_callout(doc, "GameManager.cs", [
        "Controlar score.",
        "Controlar tiempo.",
        "Contar coleccionables.",
        "Mostrar victoria.",
        "Mostrar Game Over.",
        "Desactivar el Player al perder.",
        "Permitir reiniciar con R.",
        "Actualizar la UI.",
    ])
    add_callout(doc, "MovingObstacle.cs", [
        "Mover un cubo rojo.",
        "Rebotar o alternar direccion.",
        "Detectar colision con el Player.",
        "Avisar al GameManager para perder la partida.",
    ])
    add_callout(doc, "Entrega parcial de esta fase", [
        "Scripts propuestos por el grupo.",
        "Responsabilidad de cada uno.",
        "Donde se colocan.",
        "Que necesita cada uno por Inspector.",
    ])

    add_heading(doc, "2:20 - 2:30 | Descanso", 1)
    add_body(doc, "Pausa breve antes de pasar al trabajo con IA.")

    add_heading(doc, "2:30 - 3:20 | Desarrollo con ChatGPT", 1)
    add_body(doc, "En esta fase se usa ChatGPT gratuito para pedir una solucion tecnica. La idea no es pedir directamente \"hazme el juego\", sino formular un prompt claro para que la IA proponga primero la arquitectura.")
    add_callout(doc, "Prompt inicial recomendado", """Actua como desarrollador Unity C# senior.

Estoy creando un minijuego 3D con primitivas basicas:
- capsula como Player
- esferas como coleccionables
- cubos rojos como obstaculos
- plano como suelo
- paredes hechas con cubos
- camara fija
- UI con ScoreText, TimeText y MessageText

Funcionamiento deseado:
- el jugador se mueve con teclado por el plano XZ
- recoge esferas para sumar puntos
- al recoger todas las esferas gana
- si toca un obstaculo pierde
- si el tiempo llega a 0 pierde
- al hacer Game Over, el Player debe desactivarse
- se puede reiniciar la escena con R

Restricciones:
- Unity 3D
- C#
- no usar nuevo Input System
- no usar Cinemachine
- no usar DOTween
- no usar ScriptableObjects
- mantenerlo simple
- usar Rigidbody para el Player
- usar triggers para coleccionables
- usar colisiones para obstaculos
- incluir null-checks simples
- explicar que revisar en Inspector

Primero:
1. dime que scripts son necesarios
2. explica brevemente la responsabilidad de cada script
3. indica en que GameObject debe ir cada script
4. indica que referencias deben asignarse en Inspector

Despues:
5. dame el codigo completo de cada script necesario""", fill=PROMPT_FILL, accent=DARK_BLUE, mono=True)
    add_callout(doc, "Objetivo pedagogico del prompt", [
        "Que scripts propone la IA.",
        "Si propone demasiados scripts.",
        "Si olvida alguna responsabilidad.",
        "Si respeta las restricciones.",
        "Si entiende que la camara es fija.",
        "Si usa Unity 3D y no Unity 2D.",
        "Si explica bien el Inspector.",
        "Si el codigo parece integrable.",
    ])
    add_callout(doc, "Preguntas para analizar la respuesta de ChatGPT", [
        "Propone una camara dinamica aunque la camara es fija?",
        "Propone un UIManager innecesario?",
        "Propone un GameManager demasiado grande?",
        "Usa Input System nuevo sin pedirlo?",
        "Usa Rigidbody2D por error?",
        "Donde coloca cada responsabilidad?",
        "Que referencias exige por Inspector?",
        "El codigo se entiende?",
        "Hay algo que no hayamos pedido?",
    ])
    add_heading(doc, "Reglas para trabajar con la IA", 2)
    add_bullets(doc, [
        "No copiar codigo sin leerlo.",
        "No pegar varios scripts sin comprobar que compilan.",
        "No aceptar sistemas extra si no hacen falta.",
        "No dejar que la IA cambie el objetivo del ejercicio.",
        "Pedir siempre explicacion de Inspector.",
        "Pedir cambios pequenos cuando haya errores.",
    ])

    add_heading(doc, "3:20 - 4:10 | Integracion y debugging en Unity", 1)
    add_body(doc, "En esta fase se copian los scripts, se asignan referencias en Inspector y se prueba la escena.")
    add_callout(doc, "Checklist de funcionamiento", [
        "Player se mueve.",
        "Player no gira raro.",
        "Score sube al recoger esferas.",
        "Las esferas desaparecen.",
        "El tiempo baja.",
        "Al recoger todo aparece victoria.",
        "Al llegar a 0 aparece Game Over.",
        "Al tocar obstaculo aparece Game Over.",
        "En Game Over el Player se desactiva.",
        "R reinicia la escena.",
    ])
    add_heading(doc, "Errores habituales", 2)
    add_callout(doc, "Error 1: NullReferenceException", [
        "Suele indicar que falta una referencia en Inspector.",
        "Comprobar textos de UI asignados.",
        "Comprobar referencia al GameManager.",
        "Comprobar referencia al Player.",
        "Comprobar objetos activos en escena.",
    ])
    add_callout(doc, "Error 2: OnTriggerEnter no se ejecuta", [
        "La esfera tiene Collider.",
        "La esfera tiene Is Trigger activado.",
        "El Player tiene Collider.",
        "El Player tiene Rigidbody.",
        "El Player tiene tag Player.",
    ])
    add_callout(doc, "Error 3: OnCollisionEnter no se ejecuta", [
        "El obstaculo tiene Collider normal.",
        "El Player tiene Collider.",
        "El Player tiene Rigidbody.",
        "El obstaculo no esta marcado como Trigger.",
    ])
    add_callout(doc, "Error 4: el input no funciona", [
        "Si aparece un error relacionado con UnityEngine.Input, puede que el proyecto este configurado para usar el nuevo Input System.",
        "Solucion simple: Edit > Project Settings > Player > Active Input Handling.",
        "Seleccionar Both o Input Manager (Old).",
        "Reiniciar Unity si lo pide.",
    ])
    add_callout(doc, "Error 5: la UI no actualiza", [
        "Los textos estan asignados al GameManager.",
        "Se esta usando TextMeshPro si el script usa TMP_Text.",
        "El Canvas esta activo.",
        "No hay errores de compilacion.",
    ])
    add_callout(doc, "Prompt de depuracion recomendado", """Tengo este error en Unity:
[pegar error completo]

Este es el script:
[pegar script]

Contexto de la escena:
- GameObject donde esta el script:
- componentes del Player:
- tag del Player:
- colliders:
- Rigidbody:
- referencias asignadas en Inspector:

Dime:
1. causa mas probable
2. como comprobarlo en Unity
3. cambio minimo para arreglarlo
4. que no deberia tocar""", fill=PROMPT_FILL, accent=DARK_BLUE, mono=True)
    add_callout(doc, "Idea clave de debugging", "Cuando se depura con IA, el contexto de la escena es tan importante como el codigo.")

    add_heading(doc, "4:10 - 4:40 | Comparacion de scripts", 1)
    add_body(doc, "Ahora se comparan tres cosas:")
    add_numbered(doc, [
        "La arquitectura pensada por el alumno.",
        "La arquitectura propuesta por ChatGPT.",
        "El codigo final corregido en clase.",
    ])
    add_callout(doc, "Criterios de comparacion", [
        "Compila?",
        "Funciona?",
        "Es simple?",
        "Se entiende?",
        "Respeta Unity 3D?",
        "Usa bien el Inspector?",
        "Tiene null-checks razonables?",
        "Separa responsabilidades?",
        "Evita sistemas innecesarios?",
        "El Game Over desactiva el Player?",
        "El reinicio funciona?",
    ])
    add_callout(doc, "Preguntas para debate", [
        "Que scripts propuso el alumno?",
        "Que scripts propuso la IA?",
        "La IA propuso scripts de mas?",
        "La IA olvido alguna responsabilidad?",
        "Que hizo mejor ChatGPT?",
        "Que hizo mejor el alumno?",
        "Que asumio mal la IA?",
        "Que dependia del Inspector?",
        "Que prompt mejoro mas el resultado?",
    ])
    add_heading(doc, "Conclusiones esperadas", 2)
    add_bullets(doc, [
        "ChatGPT puede acelerar mucho la escritura de scripts.",
        "La IA puede proponer arquitectura, no solo codigo.",
        "La arquitectura propuesta debe revisarse.",
        "En Unity, muchas decisiones viven fuera del script.",
        "El Inspector es parte del sistema.",
        "Un prompt con restricciones claras reduce errores.",
        "El programador sigue siendo responsable de validar.",
    ])

    add_heading(doc, "4:40 - 5:00 | Cierre y entrega final", 1)
    add_callout(doc, "Entrega minima del alumnado", [
        "Escena jugable.",
        "Scripts funcionando.",
        "Prompt usado con IA.",
        "Lista de scripts propuesta por el alumno.",
        "Lista de scripts propuesta por la IA.",
        "Comparacion breve entre ambas.",
        "Lista de errores encontrados.",
    ])
    add_callout(doc, "Cierre conceptual", "La IA no solo escribe codigo. Tambien propone una arquitectura. Nuestro trabajo como desarrolladores es evaluar si esa arquitectura es minima, clara y compatible con la escena. En Unity, la respuesta correcta vive en la relacion entre codigo, GameObjects, Inspector, fisicas, UI y prueba en Play Mode.")

    add_heading(doc, "Anexo: scripts finales esperados", 1)
    add_body(doc, "La practica puede resolverse con estos cuatro scripts:")
    add_callout(doc, "Scripts", [
        "PlayerMovement.cs",
        "Collectible.cs",
        "GameManager.cs",
        "MovingObstacle.cs",
    ])
    add_callout(doc, "PlayerMovement.cs", [
        "Responsabilidad: mover al Player con teclado usando Rigidbody.",
        "Debe ir en: Capsule / Player.",
        "Revisar en Inspector: Rigidbody, Collider y tag Player.",
    ])
    add_callout(doc, "Collectible.cs", [
        "Responsabilidad: detectar cuando el Player recoge una esfera, avisar al GameManager y desactivar la esfera.",
        "Debe ir en: cada esfera coleccionable.",
        "Revisar en Inspector: Collider con Is Trigger activado y referencia al GameManager si el script la necesita.",
    ])
    add_callout(doc, "GameManager.cs", [
        "Responsabilidad: controlar puntuacion, tiempo, coleccionables, victoria, derrota, Player y reinicio.",
        "Debe ir en: GameObject vacio llamado GameManager.",
        "Revisar en Inspector: ScoreText, TimeText, MessageText, Player, tiempo inicial y numero de coleccionables.",
    ])
    add_callout(doc, "MovingObstacle.cs", [
        "Responsabilidad: mover un obstaculo rojo, detectar colision con el Player y avisar al GameManager.",
        "Debe ir en: cada cubo rojo que actue como obstaculo.",
        "Revisar en Inspector: Collider normal, Is Trigger desactivado, referencia al GameManager, velocidad y distancia.",
    ])

    add_heading(doc, "Anexo: prompt corto para pedir mejoras", 1)
    add_callout(doc, "Prompt", """Este script funciona, pero quiero revisarlo como ejercicio docente.

Analizalo para Unity 3D y dime:
1. si cumple su responsabilidad
2. si tiene dependencias de Inspector
3. si hay riesgo de NullReferenceException
4. si esta mezclando responsabilidades
5. que cambio minimo harias para mejorarlo

No propongas sistemas nuevos.
No uses Cinemachine.
No uses nuevo Input System.
No uses DOTween.""", fill=PROMPT_FILL, accent=DARK_BLUE, mono=True)

    add_heading(doc, "Anexo: prompt corto para comparar solucion humana e IA", 1)
    add_callout(doc, "Prompt", """Tengo dos propuestas de scripts para un minijuego Unity 3D.

Propuesta del alumno:
[pegar lista]

Propuesta de ChatGPT:
[pegar lista]

El juego necesita:
- mover un Player
- recoger esferas
- sumar score
- controlar tiempo
- detectar victoria
- detectar Game Over
- mover obstaculos
- reiniciar con R
- desactivar el Player al perder

Compara ambas propuestas segun:
1. simplicidad
2. separacion de responsabilidades
3. facilidad de asignar en Inspector
4. riesgos de errores
5. adecuacion para una practica docente

Dame una conclusion breve.""", fill=PROMPT_FILL, accent=DARK_BLUE, mono=True)

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
